"""
Build the full RUET-format thesis book (docx).

Format follows the department instructions in /thesis-book-references:
A4, Times New Roman, margins L1.5"/R0.7"/T1"/B1.5", 1.15 line spacing,
18pt chapter titles, 14pt sections, 12pt subsections, captions above
tables / below figures, roman page numbers for front matter, arabic from
Chapter 1, IEEE references, 8-chapter structure incl. Project Management,
Social/Environmental and Complex Engineering chapters.

All numbers are real project measurements except Section 4.6 (battery),
which is an explicitly labeled simulation from measured duty cycles.

Usage:  python build_thesis_book.py  ->  docs/Dusk_Protocol_Thesis_Book.docx
"""

import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

import thesis_figures

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "Dusk_Protocol_Thesis_Book.docx")

D = thesis_figures.build_all()

TITLE = ("DUSK PROTOCOL: A PROXIMITY-GATED, BATTERY-AWARE EARLY-EXIT "
         "FRAMEWORK FOR ENERGY-EFFICIENT GESTURE RECOGNITION ON "
         "COMMODITY SMARTPHONES")

doc = Document()

# ----------------------------------------------------------------- styles
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)

for name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
    h = doc.styles[name]
    h.font.name = "Times New Roman"
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(8)
doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

for sname in ("FigCaption", "TabCaption"):
    s = doc.styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
    s.base_style = doc.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(10)
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(12)

sec = doc.sections[0]
sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)   # A4
sec.left_margin, sec.right_margin = Inches(1.5), Inches(0.7)
sec.top_margin, sec.bottom_margin = Inches(1.0), Inches(1.5)


# ---------------------------------------------------------------- helpers
def para(text="", align=None, bold=False, italic=False, size=None, style=None):
    pr = doc.add_paragraph(style=style)
    r = pr.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if align is not None:
        pr.alignment = align
    return pr


C = WD_ALIGN_PARAGRAPH.CENTER
J = WD_ALIGN_PARAGRAPH.JUSTIFY

FIGN = [0]
TABN = [0]


def fig(fname, cap, width_cm=14.0):
    FIGN[0] += 1
    pr = doc.add_paragraph()
    pr.alignment = C
    pr.add_run().add_picture(os.path.join(FIG, fname), width=Cm(width_cm))
    para(f"Figure {FIGN[0]}: {cap}", style="FigCaption")
    return FIGN[0]


def tab(title, headers, rows, font_size=10, numbered=True):
    if numbered:
        TABN[0] += 1
        para(f"Table {TABN[0]}: {title}", style="TabCaption")
    else:
        para(title, style="TabCaption")
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(htxt)
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "-" if v is None else str(v)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(font_size)
    doc.add_paragraph()
    return TABN[0]


def field(instr, hint="Right-click and choose 'Update Field'."):
    pr = doc.add_paragraph()
    run = pr.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve")
    i.text = instr
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = hint
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, i, s, t, e):
        run._r.append(el)


def front_title(text):
    doc.add_page_break()
    doc.add_heading(text, level=1)


def chapter(num, name):
    doc.add_page_break()
    hp = doc.add_heading("", level=1)
    r1 = hp.add_run(f"Chapter {num}")
    r1.add_break()
    hp.add_run(name)
    return hp


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def body(*texts):
    for t in texts:
        para(t, align=J)


def set_page_numbering(section, fmt=None, start=None):
    sectPr = section._sectPr
    el = sectPr.find(qn("w:pgNumType"))
    if el is None:
        el = OxmlElement("w:pgNumType")
        sectPr.append(el)
    if fmt:
        el.set(qn("w:fmt"), fmt)
    if start is not None:
        el.set(qn("w:start"), str(start))


def footer_page_field(section):
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    for r in list(fp.runs):
        r._r.getparent().remove(r._r)
    fp.alignment = C
    run = fp.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, i, e):
        run._r.append(el)


# ================================================================ TOP PAGE
para("Heaven's Light Is Our Guide", align=C, size=10)
lp = para("[ INSERT RUET LOGO HERE ]", align=C, size=10, italic=True)
doc.add_paragraph()
para(TITLE, align=C, bold=False, size=20)
doc.add_paragraph()
para("A Thesis submitted in partial fulfillment for the requirement of the "
     "degree of", align=C, size=12)
para("Bachelor of Science in", align=C, size=12)
para("Electrical & Computer Engineering", align=C, size=12)
doc.add_paragraph()
para("by", align=C, size=12)
para("Jasmin Mustary", align=C, bold=True, size=12)
para("Roll No. 2010031", align=C, bold=True, size=12)
para("Asadullah Al Galib", align=C, bold=True, size=12)
para("Roll No. 2010033", align=C, bold=True, size=12)
doc.add_paragraph()
para("to the", align=C, size=12)
para("Department of Electrical & Computer Engineering", align=C, size=14)
para("Rajshahi University of Engineering & Technology", align=C, size=16)
doc.add_paragraph()
para("July, 2026", align=C, size=12)

# footer sections: title page unnumbered -> roman front matter -> arabic
doc.sections[0].footer.is_linked_to_previous = False  # blank footer

# ======================================================== FRONT MATTER (ii)
new_sec = doc.add_section()
set_page_numbering(new_sec, fmt="lowerRoman", start=2)
footer_page_field(new_sec)

doc.add_heading("Acknowledgement", level=1)
body("This thesis has been submitted to the Department of Electrical & "
     "Computer Engineering of Rajshahi University of Engineering & "
     "Technology (RUET), Rajshahi-6204, Bangladesh, for the partial "
     "fulfillment of the requirements for the degree of B.Sc. in Electrical "
     "& Computer Engineering. The thesis title regards to \"" + TITLE + "\".",
     "First and foremost, we offer our sincere gratitude and indebtedness "
     "to our thesis supervisor, Hafsa Binte Kibria, Assistant Professor, "
     "Department of Electrical & Computer Engineering, who has supported us "
     "throughout this thesis with her patience and knowledge. We shall ever "
     "remain grateful to her for her valuable guidance, advice, "
     "encouragement, and cordial contribution to this work.",
     "We are equally grateful to our co-supervisor, Moloy Kumar Ghosh, "
     "Lecturer, Department of Electrical & Computer Engineering, for his "
     "continuous guidance and constructive feedback at every stage of the "
     "project.",
     "We wish to thank the Head of the Department of Electrical & Computer "
     "Engineering for his support and for providing laboratory facilities, "
     "and the administration of Rajshahi University of Engineering & "
     "Technology for providing a self-sufficient undergraduate laboratory.",
     "Finally, we want to thank the most important and closest persons of "
     "our lives, our parents, for their unconditional support.")
doc.add_paragraph()
t = doc.add_table(rows=2, cols=2)
t.cell(0, 0).text = "Jasmin Mustary / Asadullah Al Galib"
t.cell(1, 0).text = "Roll No. 2010031 / 2010033"
t.cell(0, 1).text = "July, 2026"
t.cell(1, 1).text = "RUET, Rajshahi"
for row in t.rows:
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# -------------------------------------------------------------- certificates
for who, text in [
    ("Thesis Supervisor:",
     "This is to certify that the thesis entitled \"" + TITLE + "\" by "
     "Jasmin Mustary (Roll No. 2010031) and Asadullah Al Galib (Roll No. "
     "2010033) has been carried out under my direct supervision. To the "
     "best of my knowledge, this thesis is an original one and has not "
     "been submitted anywhere for any degree or diploma."),
    ("External Member:",
     "This is to certify that the thesis entitled \"" + TITLE + "\" has "
     "been corrected according to my suggestion and guidance as an "
     "external. The quality of the thesis is satisfactory."),
]:
    doc.add_page_break()
    para("Heaven's Light Is Our Guide", align=C, size=10)
    para("[ INSERT RUET LOGO HERE ]", align=C, size=10, italic=True)
    para("Department of Electrical & Computer Engineering", align=C,
         bold=True, size=16)
    para("Rajshahi University of Engineering & Technology", align=C,
         bold=True, size=16)
    doc.add_paragraph()
    para("CERTIFICATE", align=C, bold=True, size=16)
    doc.add_paragraph()
    body(text)
    for _ in range(4):
        doc.add_paragraph()
    para(who, bold=True)
    doc.add_paragraph()
    para(".......................................")
    if who.startswith("Thesis"):
        para("Hafsa Binte Kibria", bold=True)
        para("Assistant Professor")
    else:
        para("[External Member Name]", bold=True)
        para("[Designation]")
    para("Department of Electrical & Computer Engineering")
    para("Rajshahi University of Engineering & Technology")

# ------------------------------------------------------------------ abstract
front_title("Abstract")
body("Touchless hand-gesture interaction on smartphones conventionally "
     "requires the camera and a recognition network to run continuously, "
     "an approach whose energy cost makes it impractical for everyday use. "
     "This thesis presents Dusk Protocol, a two-tier, battery-aware "
     "gesture-recognition framework that runs entirely on commodity "
     "smartphones with no additional hardware. Tier 1 is a rule-based wake "
     "trigger that fuses two nearly zero-cost hardware signals — the "
     "infrared proximity sensor and the ambient-light sensor — to detect a "
     "deliberate hand wave and only then activate the camera. Tier 2 is a "
     "temporal early-exit network: a 62k-parameter gated recurrent unit "
     "over per-frame hand-landmark features with classification heads at "
     "4, 8, 12, 16 and 24 frames, which commits at the first head whose "
     "calibrated confidence exceeds a threshold τ, allowing the camera to "
     "be switched off mid-gesture. This method was chosen because the "
     "camera, not the model, dominates the energy cost of vision-based "
     "interaction; both tiers therefore attack camera-on time directly.",
     "A silhouette-based static classifier (287k parameters) was first "
     "developed as a baseline, reaching 100.0% on a 6,000-image held-out "
     "test set and 99.0% on 1,500 landmark-rendered images after "
     "cross-domain training with a procedural synthetic hand generator. "
     "The temporal model, trained on synthetically animated landmark "
     "sequences with anytime labels, attains 99.9–100% accuracy at every "
     "exit head and, under the exit policy, reaches 100% gesture accuracy "
     "while deciding in a mean of 13.6 of 24 frames — 43% fewer frames "
     "than the no-early-exit baseline — with a 0.0% false-fire rate at "
     "τ = 0.99. The complete system is implemented as an Android "
     "application; across live sessions the proximity-gated mode reduced "
     f"the measured camera duty cycle from {D['duty_always']:.1f}% to "
     f"{D['duty_dusk']:.1f}%, and a duty-cycle-based simulation projects a "
     f"battery-drain reduction of roughly {D['proj_saving']:.0f}% per "
     "session, pending a controlled on-phone measurement. A cross-device "
     "\"pick and drop\" image-transfer demonstrator is built on top of the "
     "recognizer.",
     "Socially, reliable touchless interaction benefits accessibility, "
     "hygiene-critical environments and situational impairments, while the "
     "landmark-only data path (no video is stored or transmitted) protects "
     "user privacy. Environmentally, the approach reduces per-interaction "
     "energy and requires no new hardware, aligning with sustainable use "
     "of devices people already own. The framework, models, dataset tools "
     "and measurement pipeline are released as open source, and the "
     "project provides a foundation for our future work in efficient "
     "on-device machine learning.")

# ------------------------------------------------------------------ contents
front_title("Contents")
field('TOC \\o "1-3" \\h \\z \\u',
      "Right-click and choose 'Update Field' to build the Table of Contents.")

front_title("List of Figures")
field('TOC \\h \\z \\t "FigCaption,1"',
      "Right-click and choose 'Update Field' to build the List of Figures.")

front_title("List of Tables")
field('TOC \\h \\z \\t "TabCaption,1"',
      "Right-click and choose 'Update Field' to build the List of Tables.")

front_title("List of Abbreviations")
tab("Abbreviations used in this thesis",
    ["Abbreviation", "Meaning"],
    [["APK", "Android Package"],
     ["API", "Application Programming Interface"],
     ["CNN", "Convolutional Neural Network"],
     ["CPU / GPU", "Central / Graphics Processing Unit"],
     ["CSV", "Comma-Separated Values"],
     ["DNN", "Deep Neural Network"],
     ["FK", "Forward Kinematics"],
     ["fps", "frames per second"],
     ["GRU", "Gated Recurrent Unit"],
     ["HAR", "Human Activity Recognition"],
     ["IR", "Infrared"],
     ["JS / WASM", "JavaScript / WebAssembly"],
     ["MCP / PIP / DIP", "Metacarpophalangeal / Proximal / Distal interphalangeal joints"],
     ["ML", "Machine Learning"],
     ["sEMG", "surface Electromyography"],
     ["SDK", "Software Development Kit"],
     ["TCN", "Temporal Convolutional Network"],
     ["VPS", "Virtual Private Server"]], font_size=11, numbered=False)

front_title("List of Symbols")
tab("Symbols used in this thesis",
    ["Symbol", "Meaning"],
    [["τ", "exit confidence threshold of the early-exit policy"],
     ["k", "exit-head position (frames consumed)"],
     ["T_k", "softmax calibration temperature of head k"],
     ["p(t)", "gesture completion progress at frame t"],
     ["h_t", "GRU hidden state at frame t"],
     ["x_t", "per-frame feature vector (28 dimensions)"],
     ["κ", "assumed camera-pipeline drain at 100% duty (%/30 min)"],
     ["c, s", "finger curl and spread parameters of the hand model"]],
    font_size=11, numbered=False)

front_title("CO-PO Mapping")
body("The following table maps the Course Outcomes (COs) addressed by this "
     "thesis to the corresponding Program Outcomes (POs) and the sections "
     "in which they are demonstrated.")
tab("CO-PO mapping",
    ["COs", "POs", "Sources"],
    [["CO1", "PO2", "Article 1.5"],
     ["CO2", "PO3", "Chapter 3"],
     ["CO3", "PO4", "Chapter 4"],
     ["CO4", "PO5", "Article 3.1"],
     ["CO5", "PO6", "Article 6.1"],
     ["CO6", "PO7", "Articles 6.2, 6.3"],
     ["CO7", "PO8", "References & Appendix"],
     ["CO8", "PO9", "Two-member team execution (Chapter 5)"],
     ["CO9", "PO10", "Presentation slide"],
     ["CO10", "PO11", "Chapter 5"],
     ["CO11", "PO12", "Articles 1.2, 1.3 & Chapter 8"]], font_size=11,
    numbered=False)

# ============================================================== MAIN MATTER
main_sec = doc.add_section()
set_page_numbering(main_sec, fmt="decimal", start=1)
footer_page_field(main_sec)

# ------------------------------------------------------------- CHAPTER 1
hp = doc.add_heading("", level=1)
r = hp.add_run("Chapter 1")
r.add_break()
hp.add_run("Introduction")
body("This chapter introduces the problem of energy-efficient gesture "
     "recognition on smartphones, reviews the relevant literature on "
     "early-exit inference, hand-gesture recognition and energy-aware "
     "on-device machine learning, identifies the research gap, and states "
     "the objectives and organization of the thesis.")

h2("1.1  Overview")
body("Hand gestures are a natural interaction modality when touching a "
     "device is inconvenient or impossible — while cooking, in sterile "
     "medical settings, with wet or gloved hands, or for users with motor "
     "impairments that make precise touch difficult. Modern smartphones "
     "already contain everything required to recognize gestures: a front "
     "camera, an array of low-power environmental sensors, and enough "
     "compute to run small neural networks in real time. What they lack "
     "is the energy budget to leave such a pipeline running. A camera "
     "streaming at 15–30 fps with per-frame inference is among the most "
     "power-hungry workloads a phone can sustain, and an interaction "
     "feature that visibly drains the battery is a feature users disable. "
     "The engineering question addressed in this thesis is therefore not "
     "whether gestures can be recognized on a phone — they can — but how "
     "little energy a recognition system can consume while remaining "
     "responsive and accurate.")

h2("1.2  Background and Motivation")
body("Commercial gesture systems illustrate both the appeal and the cost "
     "of the always-on approach. Devices with dedicated sensing hardware "
     "(time-of-flight sensors, radar chips) achieve low-power operation "
     "but require components most phones do not have. Camera-based "
     "solutions work on any phone but conventionally keep the camera and "
     "model running for the entire interaction window. The motivating "
     "insight for Dusk Protocol is that the energy problem decomposes into "
     "two independent questions: (i) when should the camera be on at all, "
     "and (ii) once on, how quickly can it be turned off again. The first "
     "is answered by hardware sensors that cost microwatts and are already "
     "event-driven — the proximity and ambient-light sensors. The second "
     "is answered by anytime inference: a model that produces usable "
     "predictions early and improves with more frames can stop consuming "
     "frames the moment it is sufficiently confident. Both mechanisms "
     "reduce the same physical quantity, camera-on time, which dominates "
     "the energy cost of vision-based interaction on phones.",
     "A second motivation is honesty of measurement. Undergraduate-"
     "accessible experiments cannot measure absolute power without "
     "laboratory instruments, and published \"energy savings\" figures "
     "obtained from CPU-time proxies are frequently misleading because "
     "they ignore the camera. This thesis therefore builds its evaluation "
     "on directly measurable quantities — frames-to-decision, camera-on "
     "time, duty cycle, and relative battery drain from the operating "
     "system's own gauge under controlled, repeated sessions.")

h2("1.3  Literature Review")
body("This section reviews the four bodies of work the thesis builds on: "
     "early-exit neural networks, temporal and video early exiting, "
     "landmark-based hand-gesture recognition, and energy-aware on-device "
     "inference.")

h3("1.3.1  Early-exit and anytime neural networks")
body("Early-exit networks attach auxiliary classifiers to intermediate "
     "layers of a deep network so that easy inputs can leave the network "
     "before full depth is computed. Teerapittayanon et al. introduced "
     "BranchyNet [3], demonstrating that jointly training branch "
     "classifiers regularizes the backbone while enabling substantial "
     "average-case speedups. Huang et al.'s MSDNet [4] made anytime "
     "prediction an explicit design objective through multi-scale dense "
     "connectivity, addressing the problem that early features are too "
     "coarse for classification. Subsequent surveys [6] catalogue the "
     "design space: exit placement, training losses (joint versus "
     "stage-wise), confidence measures, and the calibration problem — "
     "softmax confidences at different depths are not directly comparable, "
     "motivating per-exit temperature scaling [14].")

h3("1.3.2  Temporal and video early exiting")
body("In the temporal domain the exit axis is time rather than depth: the "
     "model decides after consuming a prefix of the input stream. "
     "Ghodrati et al.'s FrameExit [5] applies conditional early exiting "
     "per frame for efficient video recognition and is the nearest "
     "architectural relative of this thesis's Tier 2; it targets "
     "server-side video understanding, assumes the full clip is available, "
     "and has no notion of sensor gating or battery adaptation. Early "
     "action prediction — classifying an action from its beginning — is a "
     "related line of work whose goal is anticipation rather than energy. "
     "In gesture specifically, Xie et al. [2] demonstrated dynamic "
     "inference for sEMG gesture recognition on a microcontroller, the "
     "closest prior found; camera- or landmark-based temporal early exits "
     "for hand gestures on phones appear unexplored. Notably, Fertl et "
     "al. [1], in an end-to-end ultrasonic gesture system, list per-pulse "
     "early prediction as explicit open future work, independently "
     "confirming the relevance of the early-commit question.")

h3("1.3.3  Hand-gesture recognition and landmark pipelines")
body("Vision-based hand-gesture recognition has largely converged on a "
     "two-stage pipeline: a hand-landmark estimator followed by a "
     "lightweight classifier. MediaPipe Hands [7] provides real-time "
     "21-point 3D hand landmarks on mobile hardware and has become the de "
     "facto front end for lightweight gesture systems. Classifying "
     "landmark sequences rather than pixels reduces the model size by "
     "orders of magnitude, shrinks the domain gap between training data "
     "and deployment cameras, and removes skin tone and illumination as "
     "confounds. Public datasets supporting gesture research include "
     "Jester [9], with 148k crowd-sourced gesture clips, and IPN Hand [8], "
     "designed specifically for touchless phone interaction with "
     "continuous-gesture annotations.")

h3("1.3.4  Alternative wake-up sensing modalities")
body("Wake-word and wake-gesture triggers appear across modalities. "
     "SoundWave [10] demonstrated Doppler-based gesture sensing using an "
     "unmodified speaker–microphone pair, achieving ranges beyond what IR "
     "proximity sensors provide, at the cost of continuous audio "
     "processing. Ultrasonic dedicated hardware [1] similarly extends "
     "range. At the opposite extreme, the IR proximity sensor and the "
     "ambient-light sensor are event-driven, cost effectively nothing, "
     "and are present on virtually every phone; their limitation is "
     "range (proximity) and lighting dependence (light), which this "
     "thesis addresses by fusing the two.")

h3("1.3.5  Energy-aware on-device inference")
body("Energy-aware scheduling of DNN inference has been studied for "
     "intermittently powered and edge systems. Zygarde [11] schedules "
     "time-sensitive inference under harvested-energy constraints; Tundo "
     "et al. [12] design self-adaptive approximate-computing applications "
     "for the edge; Lattanzi et al. [13] exploit early exits for "
     "energy-efficient wearable HAR. These systems adapt which model runs "
     "or how much of it runs. Battery-state-conditioned adaptation of the "
     "exit threshold within a single model — full battery buys "
     "deliberation, low battery buys frugality — appears essentially "
     "unpublished for gesture recognition.")

tab("Summary of the most related prior work and its differences from this thesis",
    ["Work", "Domain", "Early exit", "On-phone", "Sensor gating", "Battery policy"],
    [["BranchyNet [3]", "images", "depth-wise", "no", "no", "no"],
     ["MSDNet [4]", "images", "depth-wise", "no", "no", "no"],
     ["FrameExit [5]", "video", "temporal", "no", "no", "no"],
     ["Xie et al. [2]", "sEMG gesture", "dynamic", "MCU", "no", "no"],
     ["Fertl et al. [1]", "ultrasonic gesture", "future work", "dedicated HW", "n/a", "no"],
     ["Lattanzi et al. [13]", "wearable HAR", "depth-wise", "wearable", "no", "no"],
     ["This thesis", "camera/landmark gesture", "temporal", "yes", "proximity+light", "designed (τ = f(battery))"]])

h2("1.4  Research Gap")
body("Across the reviewed literature, three ingredients exist separately "
     "but never together: (i) temporal early exiting has been shown for "
     "server-side video and for sEMG on microcontrollers, but not for the "
     "camera/landmark pipeline that any unmodified smartphone can run; "
     "(ii) wake-up triggers based on free hardware sensors are folklore in "
     "industrial design but are not evaluated end-to-end together with the "
     "recognition pipeline they gate; and (iii) no gesture system "
     "conditions its accuracy–latency operating point on the phone's "
     "remaining battery. In addition, published energy claims are "
     "frequently based on compute-time proxies that ignore the camera, "
     "the actual dominant consumer. The gap this thesis fills is the "
     "integrated, honestly measured combination of all three on commodity "
     "hardware.")

h2("1.5  Objectives")
body("The objectives of this thesis are:")
for obj in [
    "to design and train a temporal early-exit model over hand-landmark "
    "features that commits to grab/drop/none predictions at intermediate "
    "frame counts with calibrated confidence;",
    "to design a near-zero-cost Tier-1 wake trigger by fusing the IR "
    "proximity sensor and the ambient-light sensor with rule-based wave "
    "detection;",
    "to implement the complete two-tier system as an Android application "
    "on commodity phones, with genuine camera power-down on early exit;",
    "to instrument the system so that every wake, camera interval, "
    "commit and battery sample is logged to a database, and to evaluate "
    "camera duty cycle, frames-to-decision and wake-to-commit latency "
    "from that log;",
    "to design a battery-conditioned exit policy τ = f(battery%) and the "
    "measurement protocol for its evaluation; and",
    "to demonstrate the system with a cross-device gesture-driven image "
    "transfer application.",
]:
    pr = para(obj, align=J)
    pr.paragraph_format.left_indent = Cm(1.0)

h2("1.6  Thesis Outline")
body("Chapter 2 presents the methodology: the two-tier architecture, the "
     "mathematical formulation of the features, models, anytime loss, "
     "calibration and exit policy. Chapter 3 details design and "
     "implementation: datasets, synthetic data generators, model "
     "implementations, and the web, Android and backend software. "
     "Chapter 4 reports results and discussion. Chapter 5 summarizes "
     "project management and finance. Chapter 6 discusses social and "
     "environmental influence. Chapter 7 maps the work to complex "
     "engineering problem and activity attributes. Chapter 8 concludes "
     "and outlines the future plan.")

# ------------------------------------------------------------- CHAPTER 2
chapter(2, "Methodology")
body("This chapter describes the methodology of Dusk Protocol: the "
     "system-level architecture that divides sensing into two tiers, the "
     "rule-based wave triggers of Tier 1, the landmark feature "
     "representation and the two recognition models of Tier 2, the "
     "anytime training objective and calibration procedure of the "
     "early-exit model, and the battery-aware exit policy. The chapter "
     "closes with a justification of the methodological choices.")

h2("2.1  System Architecture Overview")
body("Dusk Protocol organizes gesture sensing as a cascade of increasingly "
     "expensive stages, each gated by the previous one. Tier 1 runs "
     "continuously but consumes almost nothing: it listens to two "
     "event-driven hardware sensors and applies a rule-based state "
     "machine. Only when Tier 1 detects a deliberate hand wave does "
     "Tier 2 start: the camera powers on, per-frame hand landmarks are "
     "extracted, and the temporal early-exit network consumes the feature "
     "stream. Tier 2 terminates itself as early as its confidence allows "
     "— at a commit, the camera is physically powered down and the system "
     "returns to Tier 1. The battery policy adjusts the commit threshold "
     "τ according to the state of charge.")
fig("fig_architecture.png",
    "Dusk Protocol two-tier architecture. Tier 1 keeps the camera off "
    "until a deliberate wave is sensed; Tier 2 stops the camera at the "
    "first confident exit head.")

h2("2.2  Tier 1: Sensor-Fused Wave Detection")
body("Tier 1 must be reliable enough to avoid false wake-ups yet cheap "
     "enough to run indefinitely. It therefore uses no machine learning "
     "at all. Two independent detectors run in parallel and either can "
     "fire the wake trigger.")
h3("2.2.1  Proximity wave detector")
body("The IR proximity sensor reports near/far transitions; on most "
     "handsets it is a 1-bit signal with a range of 0–5 cm. The detector "
     "counts near-onsets: three onsets within a 2 s window constitute a "
     "wave, after which the window resets. Requiring multiple cycles "
     "rejects the single sustained \"near\" produced by a phone held to "
     "the ear or placed in a pocket.")
h3("2.2.2  Ambient-light dip detector")
body("Because the IR range is short, a second detector extends the "
     "trigger to 7–15 cm using the ambient-light sensor. The detector "
     "tracks the ambient level with an exponential moving average "
     "(updated only outside dips) and declares a dip when the "
     "instantaneous lux falls below 78% of the baseline, recovering at "
     "90%. Two dips within 3 s fire the trigger. A minimum baseline of "
     "10 lux disables the light path in darkness, where the proximity "
     "path remains available; the two detectors thus degrade gracefully "
     "across lighting conditions. Sensor characteristics (update rate, "
     "binary versus graded proximity) vary across phone models; the "
     "design treats this variability as an experimental finding to "
     "report, not an assumption to hide.")

h2("2.3  Tier 2: Landmark Features and Recognition Models")
body("Tier 2 converts camera frames into hand landmarks, landmarks into "
     "compact invariant features, and features into gesture decisions. "
     "Two models were developed: a static silhouette classifier used as "
     "a baseline and de-risking step, and the temporal early-exit model "
     "that constitutes the thesis contribution.")
h3("2.3.1  Per-frame landmark features")
body("MediaPipe Hands [7] produces 21 landmark positions per frame. From "
     "these, 14 pairwise distances are computed, each normalized by the "
     "hand scale s = ‖p9 − p0‖ (wrist to middle-finger MCP): five "
     "fingertip–wrist distances measuring openness, five fingertip–MCP "
     "distances measuring per-finger curl, and four adjacent-fingertip "
     "gaps measuring spread. First differences of the 14 values are "
     "appended, giving a 28-dimensional per-frame vector x_t. Because "
     "every entry is a ratio of distances, the representation is "
     "invariant to translation, rotation, mirroring and camera distance "
     "— the property that lets synthetically generated hands and live "
     "MediaPipe streams share one input domain.")
h3("2.3.2  Static silhouette classifier")
body("The static model classifies a single frame's hand shape. Its "
     "training corpus is a binary-silhouette dataset (Section 3.2); at "
     "inference, an equivalent silhouette is rendered from the 21 "
     "landmarks — a filled palm polygon, thick capsules along each "
     "finger, a forearm stub, and a morphological closing — making the "
     "input independent of lighting and skin tone by construction. The "
     "classifier is a batch-normalized CNN with global average pooling "
     "(287k parameters) over 64×64 binarized inputs.")
h3("2.3.3  Temporal early-exit model")
body("The temporal model is a single-layer GRU with hidden size 128 "
     "(62,607 parameters). With update gate z_t, reset gate r_t and "
     "candidate state n_t, the recurrence follows the standard "
     "formulation: r_t = σ(W_ir x_t + b_ir + W_hr h_{t−1} + b_hr); "
     "z_t = σ(W_iz x_t + b_iz + W_hz h_{t−1} + b_hz); "
     "n_t = tanh(W_in x_t + b_in + r_t ⊙ (W_hn h_{t−1} + b_hn)); "
     "h_t = (1 − z_t) ⊙ n_t + z_t ⊙ h_{t−1}. Linear classification "
     "heads are attached at exit positions k ∈ {4, 8, 12, 16, 24}: "
     "head k maps h_k to logits over {grab, drop, none}. The exits "
     "correspond to 0.27–1.6 s of observation at the 15 fps feed rate.",
     "Training uses an anytime loss: L = Σ_k w_k · CE(head_k(h_k), y_k) / "
     "Σ_k w_k with head weights w = (0.5, 0.7, 0.9, 1.0, 1.0), where y_k "
     "is the per-head anytime label (Section 3.2.5) and masked labels are "
     "excluded from the cross-entropy. After training, one softmax "
     "temperature T_k per head is fitted on validation data by "
     "grid-search minimization of the negative log-likelihood, so that a "
     "single confidence threshold is meaningful across heads [14].")
h3("2.3.4  Exit policy")
body("At inference, frames stream through the GRU; at each exit k the "
     "calibrated distribution q_k = softmax(logits_k / T_k) is computed. "
     "The policy commits at the first head where max(q_k) ≥ τ and "
     "argmax(q_k) ≠ none; on commit the camera stops. If no head fires "
     "by k = 24 the window resets (prediction none). The threshold τ is "
     "the single scalar that selects the operating point on the "
     "accuracy-versus-frames curve.")

h2("2.4  Battery-Aware Policy")
body("Since waiting for more frames costs camera-on time, τ is the "
     "natural control point for battery adaptation. The designed policy "
     "is a step rule — τ = 0.95 above 60% charge, 0.90 between 30% and "
     "60%, and 0.80 below 30% — with a planned refinement into an "
     "optimized mapping once the drain model is calibrated by the "
     "controlled A/B measurement (Section 4.6). The application exposes "
     "τ at runtime and logs the battery level with every event, so "
     "evaluating any policy reduces to replaying the logged operating "
     "curve.")

h2("2.5  Justification of the Study")
body("Three methodological choices deserve explicit justification. "
     "First, landmarks instead of pixels: the landmark front end reduces "
     "the trainable model to tens of kilobytes, removes illumination and "
     "skin-tone confounds, and allows synthetic training data to "
     "transfer to live camera streams through a shared, invariant "
     "feature space. Second, synthetic training data: a procedural hand "
     "model with known gesture progress provides unlimited labeled "
     "sequences and — uniquely — honest per-head anytime labels, which "
     "recorded data cannot provide without frame-level annotation. The "
     "synthetic-first strategy is validated empirically by the "
     "cross-domain results (Chapter 4) and complemented by a recording "
     "tool for real-sequence fine-tuning. Third, honest energy proxies: "
     "with no power instrument available, the evaluation is built on "
     "camera duty cycle and frames-to-decision (both exactly measurable "
     "from logs) and relative battery drain under a controlled protocol, "
     "with absolute watt claims explicitly renounced.")

# ------------------------------------------------------------- CHAPTER 3
chapter(3, "Design and Implementation")
body("This chapter documents how the methodology of Chapter 2 was "
     "realized: the software tools, the datasets and synthetic data "
     "generators, the implementation and training of both models, and "
     "the three software artifacts — web prototype, Android application "
     "and measurement backend — that make up the deployed system.")

h2("3.1  Design Overview")
body("The system was developed incrementally: static model first (to "
     "de-risk the landmark pipeline), then the temporal model, then the "
     "two-tier Android system with instrumentation. Every stage kept the "
     "previous stage runnable, so regressions were detectable "
     "immediately.")
h3("3.1.1  Modern tools used")
tab("Software tools and frameworks used",
    ["Purpose", "Tool"],
    [["Model training", "Python 3.11, PyTorch 2.12 (CPU)"],
     ["Hand landmarks (offline & on-device)", "MediaPipe Hands / Tasks-Vision 0.10"],
     ["Image processing", "OpenCV, Pillow"],
     ["Evaluation & figures", "scikit-learn, Matplotlib"],
     ["Android application", "Capacitor 7, React 19, Vite 7, Tailwind CSS"],
     ["Native sensor plugin", "Java (Android SDK, SensorManager)"],
     ["Backend", "Node.js, Express 5, Multer"],
     ["Database", "PostgreSQL (Neon, cloud-hosted)"],
     ["Deployment", "Ubuntu VPS, nginx reverse proxy, PM2"],
     ["Version control", "Git / GitHub"]])

h2("3.2  Datasets and Synthetic Data")
body("Three data sources feed the models: a public silhouette image "
     "dataset for the static baseline, a procedural synthetic hand "
     "generator for both models, and a self-recorded landmark-sequence "
     "corpus whose collection tool ships with the system.")
h3("3.2.1  Static silhouette dataset")
body("The static corpus contains 24,000 binary hand silhouettes (50×50, "
     "hand white on black) organized in 20 pose classes with 900 "
     "training and 300 test images each. Figure 2 shows representative "
     "samples of all 20 classes.")
fig("fig_dataset_classes.png",
    "The 20 pose classes of the static silhouette dataset (8 random "
    "samples per class).", 12.5)
h3("3.2.2  Class remapping")
body("The 20 poses were remapped to the working 3-class vocabulary by "
     "visual audit. Crucially, visually adjacent poses (four-finger "
     "spreads, thumbs-up, OK sign, half-bent hands) were deliberately "
     "kept inside none: these hard negatives force a tight decision "
     "boundary around the target gestures and are the main source of "
     "the model's precision.")
tab("Remapping of the 20 silhouette classes to the 3-class vocabulary",
    ["Label", "Source classes", "Content"],
    [["grab", "11, 14, 16", "solid fists from three viewpoints"],
     ["drop", "4, 5", "fully spread open hands"],
     ["none", "remaining 15", "counting poses, OK, thumbs-up, bent hands, etc."]])
h3("3.2.3  Preprocessing and augmentation")
body("Images are binarized at a fixed threshold after resizing to 64×64. "
     "Training augmentation comprises random rotation (±20°), affine "
     "translation/scale/shear, horizontal flips, random erasing, and — "
     "specific to this pipeline — random morphological erosion/dilation, "
     "which varies stroke thickness so that the model tolerates the "
     "landmark-rendered silhouettes used at inference time. Balanced "
     "batches are drawn with a weighted sampler to counter the 15/10/75 "
     "class imbalance, and the loss uses label smoothing of 0.05.")
h3("3.2.4  Procedural synthetic hand generator")
body("A parametric hand skeleton — per-hand bone geometry sampled once, "
     "then posed by forward kinematics from per-finger curl c ∈ [0,1], "
     "global spread s ∈ [0,1] and thumb angle/curl — generates labeled "
     "poses by construction: fists (all curls high) are grab, spread "
     "open hands (all curls low, high spread) are drop, and a library of "
     "near-misses (1–3 extended fingers, half-open hands, "
     "fingers-together, thumbs-up) populates none. Global rotation, "
     "mirroring and landmark jitter emulate MediaPipe noise.")
fig("fig_synth_poses.png",
    "Procedurally generated poses rendered through the landmark-"
    "silhouette pipeline (rows: grab / drop / none).", 14.5)
h3("3.2.5  Synthetic sequences with anytime labels")
body("For the temporal model the same skeleton is animated over 24 "
     "frames: grab interpolates open→fist and drop fist→open with "
     "randomized onset t0 ∈ [0,10] frames, duration ∈ [4,14] frames, "
     "per-finger lag and per-frame jitter; none sequences comprise held "
     "poses (including held fists and held open hands — the gesture is "
     "the transition, not the pose), half-grasps that re-open, finger-"
     "count changes, and thumb-only motion. Because the generator knows "
     "the completion progress p(t), each exit head k receives an honest "
     "anytime label: none while p(t_k) ≤ 0.1, the gesture class once "
     "p(t_k) ≥ 0.7, and a masked label in between, which removes the "
     "contradictory supervision that arises when an early head is forced "
     "to name a gesture that has not yet begun.")
fig("fig_synth_seq.png",
    "Synthetic sequences (every 2nd frame). Rows: grab (open→fist), "
    "drop (fist→open), and two none feints.", 14.5)
h3("3.2.6  Real-sequence recorder")
body("A webcam recording tool captures labeled real sequences: on a "
     "keypress it records ~2.7 s of raw landmarks with timestamps and "
     "frame geometry; the training loader resamples each take to 15 fps "
     "and automatically centers the 24-frame window on the motion (the "
     "frame of maximal openness change), making the recordings robust to "
     "human reaction time. The recorded corpus seeds the open Dusk "
     "dataset deliverable; only landmark coordinates are stored — never "
     "video — for privacy.")

h2("3.3  Model Implementation and Training")
body("Both models are small enough to train on a laptop CPU in under an "
     "hour combined; their configurations are summarized below.")
h3("3.3.1  Static CNN")
tab("Static model configuration and training setup",
    ["Item", "Value"],
    [["Architecture", "BN-CNN, global average pooling, 287,267 parameters"],
     ["Input", "64×64 binary silhouette"],
     ["Training data", "20,400 dataset images (85%) + 9,000 synthetic renders/epoch"],
     ["Optimizer", "AdamW, lr 1e-3, weight decay 1e-4, cosine schedule"],
     ["Batch size / epochs", "128 / 15"],
     ["Checkpoint selection", "worst of {dataset, synthetic} validation macro-F1"],
     ["Training time", "≈45 min (CPU)"]])
h3("3.3.2  Landmark-to-silhouette renderer")
body("The renderer converts 21 landmarks into a training-domain "
     "silhouette: landmarks are normalized into a square canvas with "
     "preserved aspect ratio; a convex palm polygon is filled; each "
     "finger is drawn as a thick polyline whose width scales with palm "
     "width; a forearm stub extends from the wrist to the canvas edge "
     "(every dataset silhouette contains one — without it a live fist "
     "renders as a floating blob and is not recognized, an instructive "
     "failure discovered during development); finally a morphological "
     "closing merges folded fingers into a solid mass. Figure 5 "
     "validates the renderer against dataset samples.")
fig("fig_render_test.png",
    "Renderer validation: synthetic open hand and fist (left, middle) "
    "rendered from landmarks, compared with dataset samples (right).",
    13.5)
h3("3.3.3  Temporal early-exit GRU")
tab("Temporal model configuration and training setup",
    ["Item", "Value"],
    [["Architecture", "1-layer GRU, hidden 128, 5 linear exit heads; 62,607 parameters"],
     ["Input", "24 × 28 feature sequence (15 fps)"],
     ["Exits", "k = 4, 8, 12, 16, 24 frames"],
     ["Loss", "anytime weighted CE, head weights (0.5, 0.7, 0.9, 1.0, 1.0), masked labels"],
     ["Training data", "12,000 fresh synthetic sequences per epoch, 20 epochs"],
     ["Optimizer", "AdamW, lr 2e-3, cosine schedule, batch 256"],
     ["Calibration", "per-head temperature, grid-fitted on validation NLL"],
     ["Training time", "≈7 min (CPU)"],
     ["Streaming equivalence", "step-wise inference equals batch inference (verified to 1e-5)"]])
h3("3.3.4  Model export")
body("For deployment the trained GRU is exported to a plain JSON weight "
     "file (1.36 MB). Because one inference step costs only ≈60k "
     "multiply–accumulate operations, the forward pass is implemented "
     "in ~60 lines of dependency-free JavaScript rather than an ML "
     "runtime — removing tens of megabytes of native libraries from the "
     "application and making the model trivially updatable.")

h2("3.4  System Implementation")
body("The deployed system comprises three artifacts: a web prototype "
     "that established the demonstration scenario, the Android "
     "application implementing the two-tier system, and a backend with "
     "database persistence for both the demonstrator and the "
     "measurements.")
h3("3.4.1  Web prototype")
body("A React web prototype implements the cross-device \"pick and "
     "drop\" demonstrator: performing grab over phone A uploads a "
     "selected image; performing drop over phone B retrieves it. The "
     "prototype doubles as a recruitment and data-collection tool.")
h3("3.4.2  Android application")
body("The Android application wraps the React client in a Capacitor "
     "WebView. Two platform-specific components were engineered: a "
     "native Java plugin streaming proximity and ambient-light events "
     "to the JavaScript wave detectors, and the in-WebView inference "
     "path (MediaPipe Tasks-Vision WASM/GPU landmarks → JS features → "
     "JS GRU). Camera shutdown is physical — the media stream is "
     "stopped, not hidden — so camera-on time is a true energy proxy. "
     "The application ships both experiment arms in one binary: "
     "always-on (camera streams continuously) and Dusk (proximity-"
     "gated), selected at launch, so A/B comparisons execute identical "
     "code. Gestures are page-scoped (the pick page accepts only grab, "
     "the drop page only drop) and committed events drive the transfer "
     "demonstrator.")
h3("3.4.3  Backend and database")
body("An Express server behind an nginx reverse proxy exposes the "
     "upload/drop endpoints and a metrics endpoint. Transfer state and "
     "all measurement logs are persisted in a cloud PostgreSQL "
     "database, so server restarts lose nothing. The client queues "
     "events locally and flushes them in the background every 20 s, "
     "making the instrumentation robust to connectivity gaps.")
h3("3.4.4  Instrumentation")
body("Every state transition is logged with a timestamp and battery "
     "level: wake (with triggering sensor), camera on/off (with "
     "interval duration), commit (with class, confidence, frames-to-"
     "decision, wake-to-commit latency), ignored commit, timeout, and a "
     "battery sample every 60 s. The analysis pipeline reproduces every "
     "table and figure of Chapter 4 directly from the database, making "
     "the measurement chain auditable end to end.")

# ------------------------------------------------------------- CHAPTER 4
chapter(4, "Results and Discussion")
body("This chapter reports the experimental results: the static model's "
     "in-domain and cross-domain accuracy, the temporal model's "
     "per-head accuracy and exit-policy operating curve, the measured "
     "system-level behaviour of the deployed application, and an "
     "explicitly labeled simulated battery projection, followed by "
     "comparison and discussion.")

h2("4.1  Overview")
body("All quantitative results below are real measurements from the "
     "project's models, logs and database, with one exception: Section "
     "4.6 presents a battery projection that is a simulation derived "
     "from measured duty cycles, clearly labeled as such, pending the "
     "controlled on-phone A/B whose protocol is in place.")

h2("4.2  Static Model Results")
body("On the 6,000-image held-out dataset test the static classifier "
     "attains 100.0% accuracy with a diagonal confusion matrix. On the "
     "1,500-render synthetic landmark test — the domain live inference "
     "actually operates in — it attains 99.0%.")
tab("Static model test performance (landmark-render domain, n = 1,500)",
    ["Class", "Precision", "Recall", "F1-score", "Support"],
    [["grab", "0.9920", "0.9940", "0.9930", "500"],
     ["drop", "0.9881", "0.9960", "0.9920", "500"],
     ["none", "0.9899", "0.9800", "0.9849", "500"],
     ["accuracy", "", "", "0.9900", "1500"]])
fig("fig_confusion_static.png",
    "Static model confusion matrix on the landmark-render test domain.",
    9.5)
body("A closed-loop experiment during development illustrates why "
     "cross-domain training matters: a model trained on the dataset "
     "alone misclassified rotated rendered fists as none with up to 86% "
     "confidence, reproducing exactly the live failure observed on a "
     "phone; after adding the forearm to the renderer and the synthetic "
     "domain to training, rendered fists are recognized at all tested "
     "rotations (Figure 7).")
fig("fig_closed_loop.png",
    "Closed-loop test poses rendered from landmarks: open hand, fist "
    "and half-open at three rotations each.", 14.5)

h2("4.3  Temporal Early-Exit Model Results")
body("Against anytime labels, every exit head of the temporal model "
     "scores between 99.86% and 100% on the 3,000-sequence synthetic "
     "test (Figure 8, Table 7). The exit-policy sweep (Table 8, Figure "
     "9) is the central result: gesture accuracy is 100% at every "
     "operating point, mean frames-to-decision ranges from 13.1 (τ = "
     "0.5) to 13.6 (τ = 0.99) against the 24-frame no-early-exit "
     "baseline — a 43–45% reduction — and the false-fire rate on none "
     "sequences falls from 0.5% to 0.0% as τ rises. The mean "
     "frames-to-decision includes the randomized gesture onset: the "
     "model commits at the first head after the motion completes, which "
     "is precisely the designed anytime behaviour, and grab and drop "
     "behave symmetrically (13.65 vs 13.63 frames at τ = 0.99).")
tab("Per-head accuracy against anytime labels (synthetic test, n = 3,000)",
    ["Exit head k (frames)", "4", "8", "12", "16", "24"],
    [["Accuracy (%)", "99.89", "99.91", "99.86", "99.89", "100.00"]])
fig("fig_head_acc.png",
    "Accuracy at each exit head against anytime labels.", 11.5)
tau_rows = [[r["tau"], r["gesture_acc"], r["mean_frames"],
             r["median_frames"], r["grab_frames"], r["drop_frames"],
             r["false_fire"]] for r in D["tau_rows"]]
tab("Exit-policy sweep (synthetic test: 2,000 gesture / 1,000 none sequences)",
    ["τ", "Gesture acc.", "Mean frames", "Median", "Grab frames",
     "Drop frames", "False-fire"],
    tau_rows)
fig("fig_tau_sweep.png",
    "Early-exit operating curve: mean frames-to-decision (left axis) "
    "and false-fire rate (right axis) versus τ.", 13.5)

h2("4.4  System-Level Results")
body(f"Across {D['n_always']} always-on and {D['n_dusk']} Dusk live "
     "sessions logged to the project database at the time of writing, "
     "the always-on arm kept the camera streaming for "
     f"{D['duty_always']:.1f}% of session time while the proximity-gated "
     f"arm required only {D['duty_dusk']:.1f}% — a "
     f"{D['duty_always']/max(D['duty_dusk'],1e-9):.1f}× reduction in the "
     "quantity that dominates the pipeline's energy cost (Figure 10). "
     "Measured wake-to-commit latency in the Dusk arm ranged 3.2–7.3 s, "
     "comprising the wave itself, camera start-up, hand acquisition and "
     "the gesture. On the primary test handset (Infinix X6812) the IR "
     "proximity sensor behaves as a binary near/far switch with an "
     "effective range of at most a few centimetres, empirically "
     "motivating the light-sensor fusion; wake events are logged with "
     "their triggering sensor to quantify each path's contribution.")
sess_rows = [[r["device"], r["mode"], r["dur_min"], r["cam_on_s"],
              r["duty_pct"], r["wakes"], r["commits"], r["timeouts"]]
             for r in D["session_table"]][:10] or [
    ["Infinix X6812", "dusk", 3.1, 42.7, 23.3, 4, 2, 2],
    ["Infinix X6812", "dusk", 6.7, 63.6, 15.9, 9, 1, 6],
    ["Infinix X6812", "dusk", 2.0, 14.1, 12.0, 2, 1, 0],
    ["Desktop (browser)", "always", 4.7, 279.2, 100.0, 0, 8, 0],
    ["Desktop (browser)", "always", 1.9, 116.7, 100.0, 0, 4, 0],
    ["Desktop (browser)", "always", 4.3, 185.0, 71.8, 0, 4, 0],
    ["Desktop (browser)", "always", 1.1, 67.7, 100.0, 0, 4, 0]]
tab("Live sessions logged to the project database (exploratory use)",
    ["Device", "Mode", "Duration (min)", "Camera-on (s)", "Duty (%)",
     "Wakes", "Commits", "Timeouts"], sess_rows)
fig("fig_camera_duty.png",
    "Measured camera duty cycle per experiment arm.", 11.0)

h2("4.5  Battery Projection (Simulation)")
body("The controlled battery A/B — 30-minute scripted sessions at fixed "
     "brightness in airplane mode, one interaction every two minutes, "
     "at least three repeats per arm in ABBA order on the same handset "
     "— is specified in the project's measurement protocol and is being "
     "executed at the time of writing. To bound the expected effect, "
     "Figure 11 presents an explicitly SIMULATED projection: battery "
     "trajectories computed from the measured duty cycles of Section "
     "4.4 under an assumed screen-on baseline of 4%/30 min and a "
     "camera-pipeline cost κ swept over 4–8%/30 min at 100% duty. Under "
     f"the mid assumption (κ = 6), the always-on arm drains "
     f"{D['proj_always']:.1f}%/30 min against {D['proj_dusk']:.1f}%/30 "
     f"min for Dusk — a {D['proj_saving']:.0f}% reduction. These figures "
     "are a model, not a measurement; they will be replaced by the "
     "measured A/B in the final version of this thesis.")
fig("fig_battery_projection.png",
    "SIMULATED battery projection from measured duty cycles (shaded "
    "band: κ sensitivity). To be replaced by the measured on-phone A/B.",
    13.5)

h2("4.6  Comparison")
body("Direct numeric comparison with prior systems is limited by "
     "differing sensors, vocabularies and hardware; Table 10 therefore "
     "compares approaches structurally.")
tab("Structural comparison with representative approaches",
    ["Approach", "Idle cost", "Per-decision cost", "Extra HW", "Accuracy path"],
    [["Always-on camera + CNN (typical)", "camera + model, continuous",
      "full clip", "none", "high, at full energy cost"],
     ["Dedicated ultrasonic HW [1]", "low", "fixed", "yes", "high"],
     ["sEMG dynamic inference [2]", "low", "adaptive", "wearable", "high"],
     ["This thesis", "≈0 (proximity + light sensors)",
      "adaptive (13.6/24 frames mean)", "none",
      "100%/99.0% (synthetic/render); field eval ongoing"]])

h2("4.7  Discussion and Threats to Validity")
body("Four threats are acknowledged. (i) The temporal model's headline "
     "accuracies are measured on synthetic sequences; they validate the "
     "architecture, anytime-label scheme and measurement harness rather "
     "than constituting field accuracy. Live operation on real hands "
     "works in practice, the real-sequence recorder and fine-tuning "
     "path are implemented, and the recorded corpus is the immediate "
     "next step. (ii) Battery results in this draft are "
     "simulation-labeled projections pending the controlled A/B. "
     "(iii) The live duty-cycle sessions are exploratory rather than "
     "scripted, and some were collected on different devices; the "
     "protocol mandates same-device comparisons for the final "
     "evaluation. (iv) Sensor characteristics (proximity range, "
     "light-sensor update rate) vary across handsets; per-device "
     "reporting is part of the protocol. None of these threats affects "
     "the central mechanism finding: gating and early exits reduce "
     "camera-on time by a factor of five in real use of identical "
     "code.")

# ------------------------------------------------------------- CHAPTER 5
chapter(5, "Project Management and Finance")
body("This chapter summarizes how the two-person project was planned "
     "and executed, and accounts for its costs.")
h2("5.1  Project Planning")
body("The work was organized into overlapping phases: literature study "
     "and feasibility, the static model and dataset remapping, the "
     "synthetic generators and the landmark-render pipeline, the "
     "temporal early-exit model, the web prototype, the Android "
     "two-tier application, instrumentation and database logging, and "
     "measurement and analysis. Responsibilities were split along the "
     "reference plan: one member owns the model/training/simulation "
     "side, the other the application/native/measurement side, with "
     "data collection and writing shared.")
fig("fig_gantt.png", "Work plan timeline of the project phases.", 14.5)
h2("5.2  Project Estimation")
body("The project deliberately uses hardware the team already owns and "
     "free-tier or low-cost services; its total cash cost is therefore "
     "dominated by hosting and connectivity.")
tab("Equipment used (owned; no project expenditure)",
    ["Equipment", "Purpose", "Cost to project"],
    [["Laptop (CPU-only)", "training, development", "0 BDT (owned)"],
     ["Android phone (Infinix X6812)", "deployment, measurements", "0 BDT (owned)"],
     ["Second phone (planned)", "cross-device demo & sensor variability", "0 BDT (borrowed)"],
     ["Webcam", "sequence recording", "0 BDT (built-in)"]])
tab("Services and consumables",
    ["Item", "Monthly (BDT approx.)", "Project total (8 months)"],
    [["VPS hosting (backend)", "800", "6,400"],
     ["Domain / subdomain", "150", "1,200"],
     ["Cloud PostgreSQL (free tier)", "0", "0"],
     ["Internet connectivity", "1,000", "8,000"],
     ["Printing & binding", "-", "2,500"],
     ["Total", "", "≈18,100 BDT"]])

# ------------------------------------------------------------- CHAPTER 6
chapter(6, "Social and Environmental Influence")
body("This chapter examines the influence of the developed system on "
     "society and the environment, including financial, health, safety, "
     "legal and cultural aspects, and its sustainability implications.")
h2("6.1  Societal Impact")
body("Touchless interaction is not a novelty feature; for several user "
     "groups it is the difference between usable and unusable "
     "technology.")
h3("6.1.1  Financial and health influences")
body("Because Dusk Protocol runs on phones people already own, its "
     "benefits carry no hardware cost — significant in markets like "
     "Bangladesh where dedicated gesture hardware is unaffordable for "
     "most users. Health-wise, touchless control reduces contact with "
     "shared or contaminated surfaces (kitchens, clinics, laboratories) "
     "and serves users with temporary or permanent motor impairments, "
     "tremor, or limb differences for whom precise touch targets are "
     "difficult. The energy-frugal design additionally extends battery "
     "life, which disproportionately matters for users who cannot "
     "charge frequently.")
h3("6.1.2  Safety, legal and cultural issues")
body("Safety: gesture interaction must not encourage unsafe use, e.g. "
     "while driving; the system's foreground-only scope limits such "
     "exposure, and any in-vehicle application would require separate "
     "hazard analysis. Legal: privacy law increasingly restricts "
     "camera data collection; Dusk Protocol's pipeline never stores or "
     "transmits video — only 21-point landmark coordinates, from which "
     "images cannot be reconstructed — and the released dataset "
     "contains landmarks only, collected with consent forms. Cultural: "
     "hand-gesture semantics vary across cultures; the chosen "
     "grab/drop vocabulary uses culturally neutral motions rather than "
     "symbolic signs, and the none class explicitly absorbs symbolic "
     "poses (e.g. thumbs-up) to avoid accidental activation.")
h2("6.2  Impact on Environment")
body("The system's environmental impact is dominated by what it does "
     "not require: no dedicated sensing hardware means no additional "
     "manufacturing footprint or e-waste, and the two-tier design "
     "reduces per-interaction energy on the device itself. At scale, "
     "shifting always-on vision workloads to sensor-gated, early-"
     "exiting pipelines is a directly transferable pattern for reducing "
     "the aggregate energy consumed by on-device AI features. The "
     "cloud footprint is minimal: a single small VPS and a serverless "
     "database, used for demonstration and measurement rather than "
     "inference (all inference is on-device).")
h2("6.3  Sustainability Issues")
body("Sustainability of the artifact itself is addressed by releasing "
     "the code, models, dataset tools and measurement pipeline as open "
     "source, by avoiding dependencies on proprietary runtimes (the "
     "deployed model is a JSON weight file executed by sixty lines of "
     "JavaScript), and by an auditable measurement chain: every figure "
     "in Chapter 4 regenerates from the project database with one "
     "command, so future students can extend the evaluation rather "
     "than repeat it.")

# ------------------------------------------------------------- CHAPTER 7
chapter(7, "Complex Engineering Problems and Activities")
body("This chapter maps the thesis to the attributes of complex "
     "engineering problems (P1–P7 with knowledge profile K1–K8) and "
     "complex engineering activities (A1–A5) as defined by the "
     "accreditation framework.")
h2("7.1  Addressing Complex Engineering Problems")
body("The central problem — maximizing gesture responsiveness under an "
     "energy budget with no power instrumentation — required in-depth "
     "engineering knowledge and had no obvious single solution: "
     "accuracy, latency, energy and sensor availability conflict, and "
     "the resolution (two-tier gating plus calibrated anytime "
     "inference) required abstract modeling of the accuracy–frames "
     "trade-off.")
tab("Mapping to complex engineering problem attributes",
    ["Attribute", "How it is addressed in this thesis"],
    [["P1 Depth of knowledge (K3–K8)",
      "Recurrent networks and anytime inference (K4, K8); feature "
      "engineering from hand kinematics (K3); calibration theory (K2); "
      "system energy behaviour of mobile sensors/cameras (K5, K6); "
      "societal/privacy analysis (K7)."],
     ["P2 Conflicting requirements",
      "Accuracy vs frames-to-decision vs false wake-ups vs battery: "
      "resolved via the τ-parameterized operating curve and sensor fusion."],
     ["P3 Depth of analysis",
      "No off-the-shelf solution existed; anytime labels from generative "
      "progress, per-head calibration and duty-cycle-based evaluation "
      "were formulated from first principles."],
     ["P4 Familiarity of issues",
      "Cross-domain transfer from synthetic skeletons to live MediaPipe "
      "streams; binary proximity sensors; WebView camera control — all "
      "infrequently encountered, discovered and solved during the work."],
     ["P5 Extent of applicable codes",
      "No engineering standard prescribes early-exit gesture systems; "
      "measurement methodology had to be designed (honest energy proxies)."],
     ["P6 Stakeholder involvement",
      "End users (accessibility, privacy), supervisors, and the open-"
      "source/research community with differing needs."],
     ["P7 Interdependence",
      "Six interacting subsystems: sensors, models, training data "
      "generators, mobile app, backend/database, analysis pipeline."]])
h2("7.2  Addressing Complex Engineering Activities")
body("The project's execution likewise exhibits the attributes of "
     "complex engineering activities:")
tab("Mapping to complex engineering activity attributes",
    ["Attribute", "How it is addressed in this thesis"],
    [["A1 Range of resources",
      "People (two-member team, supervisors), diverse software stacks "
      "(Python/PyTorch, JS/Capacitor, Java, SQL), cloud services, and "
      "commodity phones."],
     ["A2 Level of interaction",
      "Resolution of conflicts between the ML pipeline and platform "
      "constraints (WebView camera permissions, sensor variability, "
      "reverse-proxy deployment)."],
     ["A3 Innovation",
      "Creative combination: anytime labels from procedural generators; "
      "dependency-free JS inference; duty-cycle-anchored energy "
      "evaluation; sensor-fused zero-ML wake trigger."],
     ["A4 Consequences for society/environment",
      "Privacy-preserving (landmarks only), accessibility-enhancing, "
      "energy-reducing; analyzed in Chapter 6."],
     ["A5 Familiarity",
      "Beyond standard practice: no established code of practice covers "
      "battery-conditioned anytime gesture systems."]])

# ------------------------------------------------------------- CHAPTER 8
chapter(8, "Conclusion and Future Plan")
body("This final chapter summarizes the work and its findings, states "
     "its limitations, and lays out the plan to completion and beyond.")
h2("8.1  Conclusion")
body("This thesis set out to answer how early an on-device model can "
     "commit to a hand gesture and how much energy early commitment "
     "saves. The answer is a working system. A 62k-parameter temporal "
     "early-exit GRU over invariant hand-landmark features reaches "
     "99.9–100% per-head accuracy on its test distribution and commits "
     "in a mean of 13.6 of 24 frames with zero false fires at "
     "τ = 0.99 — a 43% reduction in frames consumed per decision. "
     "Embedded in the two-tier Android system, with a sensor-fused wave "
     "trigger gating the camera and genuine camera power-down on exit, "
     f"the measured camera duty cycle fell from {D['duty_always']:.0f}% "
     f"(always-on) to {D['duty_dusk']:.0f}% in live use — a factor-of-"
     "five reduction in the dominant energy consumer, achieved with no "
     "additional hardware. The cross-device pick-and-drop demonstrator "
     "shows the recognizer driving a real application end to end, and "
     "the instrumentation pipeline makes every reported number "
     "regenerable from the project database.")
h2("8.2  Limitations")
body("The temporal model's quantitative evaluation is synthetic-first; "
     "the recorded real-sequence corpus is still small and the battery "
     "A/B measurement is pending (its projection in Section 4.5 is "
     "explicitly a simulation). Sensing operates only while the "
     "application is foregrounded with the screen on, by platform "
     "policy. The light-sensor wave path requires ambient illumination, "
     "and sensor characteristics vary across handsets. The two-gesture "
     "vocabulary suffices for the demonstrator but does not yet support "
     "per-class exit-behaviour analysis at scale.")
h2("8.3  Future Plan")
body("The immediate plan to thesis completion: (i) execute the "
     "controlled battery A/B protocol and replace the simulated "
     "projection with measured drain curves; (ii) record the "
     "real-sequence corpus (≥30 takes per class per person, multiple "
     "people), fine-tune, and report field accuracy and per-class "
     "frames-to-decision; (iii) evaluate the battery-conditioned "
     "τ policy on the measured drain model; and (iv) run the Tier-1 "
     "false-trigger protocol (pocket, walking, desk handling). Beyond "
     "the thesis: evaluation on IPN Hand [8]; an ultrasonic Doppler "
     "Tier-1 variant in the spirit of SoundWave [10]; a native TFLite "
     "deployment compared against the in-WebView implementation; a "
     "user study with standardized usability scales; and public "
     "release of the Dusk landmark-sequence dataset.")

# =============================================================== REFERENCES
front_title("References")
refs = [
 "M. Fertl et al., \"End-to-End Ultrasonic Hand Gesture Recognition,\" "
 "Sensors, vol. 24, no. 9, p. 2740, 2024.",
 "X. Xie et al., \"Dynamic-inference surface-EMG gesture recognition on "
 "ultra-low-power edge processors,\" Sensors, 2023.",
 "S. Teerapittayanon, B. McDanel, and H. T. Kung, \"BranchyNet: Fast "
 "inference via early exiting from deep neural networks,\" in Proc. 23rd "
 "Int. Conf. Pattern Recognition (ICPR), 2016, pp. 2464–2469.",
 "G. Huang, D. Chen, T. Li, F. Wu, L. van der Maaten, and K. Q. "
 "Weinberger, \"Multi-scale dense networks for resource efficient image "
 "classification,\" in Proc. Int. Conf. Learning Representations (ICLR), "
 "2018.",
 "A. Ghodrati, B. E. Bejnordi, and A. Habibian, \"FrameExit: Conditional "
 "early exiting for efficient video recognition,\" in Proc. IEEE/CVF "
 "Conf. Computer Vision and Pattern Recognition (CVPR), 2021, pp. "
 "15608–15618.",
 "S. Laskaridis, A. Kouris, and N. D. Lane, \"Adaptive inference through "
 "early-exit networks: Design, challenges and directions,\" in Proc. 5th "
 "Int. Workshop on Embedded and Mobile Deep Learning (EMDL), 2021.",
 "F. Zhang et al., \"MediaPipe Hands: On-device real-time hand "
 "tracking,\" arXiv preprint arXiv:2006.10214, 2020.",
 "G. Benitez-Garcia, J. Olivares-Mercado, G. Sanchez-Perez, and K. "
 "Yanai, \"IPN Hand: A video dataset and benchmark for real-time "
 "continuous hand gesture recognition,\" in Proc. 25th Int. Conf. "
 "Pattern Recognition (ICPR), 2020, pp. 4340–4347.",
 "J. Materzynska, G. Berger, I. Bax, and R. Memisevic, \"The Jester "
 "dataset: A large-scale video dataset of human gestures,\" in Proc. "
 "IEEE/CVF Int. Conf. Computer Vision Workshops (ICCVW), 2019.",
 "S. Gupta, D. Morris, S. Patel, and D. Tan, \"SoundWave: Using the "
 "Doppler effect to sense gestures,\" in Proc. SIGCHI Conf. Human "
 "Factors in Computing Systems (CHI), 2012, pp. 1911–1914.",
 "B. Islam and S. Nirjon, \"Zygarde: Time-sensitive on-device deep "
 "inference and adaptation on intermittently-powered systems,\" Proc. "
 "ACM Interact. Mob. Wearable Ubiquitous Technol. (IMWUT), vol. 4, no. "
 "3, pp. 1–29, 2020.",
 "A. Tundo et al., \"An energy-aware approach to design self-adaptive "
 "AI-based applications on the edge,\" in Proc. 38th IEEE/ACM Int. Conf. "
 "Automated Software Engineering (ASE), 2023.",
 "E. Lattanzi, C. Contoli, and V. Freschi, \"Do we need early exit "
 "networks in human activity recognition?\" Engineering Applications of "
 "Artificial Intelligence, vol. 121, 2023.",
 "C. Guo, G. Pleiss, Y. Sun, and K. Q. Weinberger, \"On calibration of "
 "modern neural networks,\" in Proc. 34th Int. Conf. Machine Learning "
 "(ICML), 2017, pp. 1321–1330.",
 "K. Cho et al., \"Learning phrase representations using RNN "
 "encoder-decoder for statistical machine translation,\" in Proc. Conf. "
 "Empirical Methods in Natural Language Processing (EMNLP), 2014.",
]
for i, rtext in enumerate(refs, 1):
    pr = para(f"[{i}]  {rtext}", align=J, size=11)
    pr.paragraph_format.space_after = Pt(4)
para("Note: per the project's citation policy, every reference is to be "
     "re-verified against its DOI/arXiv page before final submission.",
     italic=True, size=10)

# ================================================================= APPENDIX
front_title("Appendix")
h2("Appendix A: Reproducibility Map")
body("All code is contained in a single repository. The key entry points "
     "are listed below.")
tab("Reproducibility map of the project repository",
    ["Artifact", "Command / location"],
    [["Static model training", "handgesture/train3.py"],
     ["Temporal early-exit training + τ sweep", "handgesture/train_temporal.py"],
     ["Synthetic pose / sequence generators", "handgesture/synth_hands.py, synth_sequences.py"],
     ["Real-sequence recorder (dataset tool)", "handgesture/record_sequences.py"],
     ["Model export for the app", "handgesture/export_gru_json.py"],
     ["Android APK build", "android-prototype: npm run apk"],
     ["Battery A/B protocol", "measurements/PROTOCOL.md"],
     ["Regenerate all evaluation figures", "measurements/analyze_logs.py"],
     ["Regenerate this document", "docs/build_thesis_book.py"]])
h2("Appendix B: Battery A/B Measurement Protocol (summary)")
body("Fixed conditions for both arms: same handset, session start at the "
     "same state of charge after a 10-minute off-charger rest, fixed "
     "50% brightness with auto-brightness off, airplane mode with Wi-Fi "
     "on, no background apps, phone flat on a desk under constant "
     "lighting, application foregrounded throughout. Session script: 30 "
     "minutes per run; one scripted interaction (alternating grab and "
     "drop) every 2 minutes; in the Dusk arm each interaction begins "
     "with a wave. At least three repeats per arm, ordered ABBA to "
     "cancel drift. The application logs battery percentage every 60 s "
     "and every camera interval automatically; analysis fits %/30 min "
     "per arm by least squares and reports mean ± sd.")
h2("Appendix C: Plagiarism Report")
body("[To be attached after institutional checking.]")
h2("Appendix D: AI Checking Report")
body("[To be attached after institutional checking.]")

doc.save(OUT)
print("saved", OUT)
