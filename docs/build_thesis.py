"""
Build the Dusk Protocol thesis draft as a .docx.

All quantitative claims are real measurements from this project except the
battery projection, which is explicitly labeled as a simulation derived
from measured camera duty cycles (to be replaced by the measured A/B —
see measurements/PROTOCOL.md).

Usage:  python build_thesis.py   ->  docs/Dusk_Protocol_Thesis_Draft.docx
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import thesis_figures

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "Dusk_Protocol_Thesis_Draft.docx")

D = thesis_figures.build_all()

doc = Document()

# ------------------------------------------------------------------ styles
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.3

for name, size, bold in [("Heading 1", 16, True), ("Heading 2", 13, True),
                         ("Heading 3", 12, True)]:
    h = doc.styles[name]
    h.font.name = "Times New Roman"
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = RGBColor(0, 0, 0)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# page numbers in footer
footer_p = doc.sections[0].footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_p.add_run()
for el, attrs, text in [("w:fldChar", {"w:fldCharType": "begin"}, None),
                        ("w:instrText", {"xml:space": "preserve"}, "PAGE"),
                        ("w:fldChar", {"w:fldCharType": "end"}, None)]:
    e = OxmlElement(el)
    for k, v in attrs.items():
        e.set(qn(k), v)
    if text:
        e.text = text
    run._r.append(e)


def p(text, style=None, align=None, italic=False, bold=False, size=None):
    para = doc.add_paragraph(style=style)
    r = para.add_run(text)
    r.italic = italic
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if align:
        para.alignment = align
    return para


def h1(text):
    doc.add_page_break()
    return doc.add_heading(text, level=1)


def h2(text):
    return doc.add_heading(text, level=2)


def h3(text):
    return doc.add_heading(text, level=3)


def caption(text):
    para = p(text, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)
    para.paragraph_format.space_after = Pt(14)
    return para


def figure(fname, caption_text, width_cm=14.5):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(os.path.join(FIG, fname), width=Cm(width_cm))
    caption(caption_text)


def table(headers, rows, caption_text=None, font_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = htxt
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(font_size)
    doc.add_paragraph()
    if caption_text:
        caption(caption_text)
    return t


def toc_field():
    para = doc.add_paragraph()
    run = para.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "Right-click here and choose 'Update Field' to build the Table of Contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for e in (begin, instr, sep, hint, end):
        run._r.append(e)


# ================================================================ TITLE PAGE
for _ in range(4):
    doc.add_paragraph()
p("DUSK PROTOCOL", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
p("A Proximity-Gated, Battery-Aware Early-Exit Framework for "
  "Energy-Efficient Gesture Recognition on Commodity Smartphones",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=15)
doc.add_paragraph()
p("A thesis submitted in partial fulfilment of the requirements for the "
  "degree of Bachelor of Science in Electrical & Computer Engineering",
  align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11)
for _ in range(2):
    doc.add_paragraph()
p("Jasmin Mustary (2010031)", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
p("Asadullah Al Galib (2010033)", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
doc.add_paragraph()
p("Supervised by", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
p("Hafsa Binte Kibria, Assistant Professor, Dept. of ECE",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
p("Moloy Kumar Ghosh, Lecturer, Dept. of ECE",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
for _ in range(2):
    doc.add_paragraph()
p("Department of Electrical & Computer Engineering",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
p("Rajshahi University of Engineering & Technology",
  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
p("July 2026 — DRAFT FOR EARLY SUBMISSION",
  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)

# ================================================================== ABSTRACT
h1("Abstract")
p("Touchless hand-gesture interaction on smartphones conventionally requires "
  "the camera and a neural network to run continuously, an approach whose "
  "energy cost makes it impractical for everyday use. This thesis presents "
  "Dusk Protocol, a two-tier, battery-aware gesture recognition system that "
  "runs entirely on commodity smartphones with no additional hardware. "
  "Tier 1 is a rule-based wake trigger that fuses two nearly zero-cost "
  "hardware signals — the infrared proximity sensor and the ambient light "
  "sensor — to detect a deliberate hand wave and only then activate the "
  "camera. Tier 2 is a temporal early-exit network: a 62k-parameter gated "
  "recurrent unit over per-frame hand-landmark features with classification "
  "heads at 4, 8, 12, 16 and 24 frames, which commits to a prediction at the "
  "first head whose calibrated confidence exceeds a threshold τ, allowing "
  "the camera to be switched off mid-gesture.")
p("A silhouette-based static classifier (287k parameters) was first developed "
  "as a baseline, reaching 100.0% on a 6,000-image held-out test set and "
  "99.0% on 1,500 landmark-rendered images after cross-domain training with "
  "a procedural synthetic hand generator. The temporal model, trained on "
  "synthetically animated landmark sequences with anytime labels, attains "
  "99.9–100% accuracy at every exit head and, under the exit policy, "
  "reaches 100% gesture accuracy while deciding in a mean of 13.6 of 24 "
  "frames (43% fewer frames than the no-early-exit baseline) with a 0.0% "
  "false-fire rate at τ = 0.99. The full system is implemented as an "
  "Android application whose recognition events, camera intervals and "
  "battery samples are logged to a cloud database. Across live sessions, "
  f"the proximity-gated mode reduced the camera duty cycle from "
  f"{D['duty_always']:.1f}% (always-on) to {D['duty_dusk']:.1f}%, and a "
  "duty-cycle-based simulation projects a corresponding battery-drain "
  f"reduction of roughly {D['proj_saving']:.0f}% per session; a controlled "
  "on-phone battery A/B measurement is in progress. A Huawei-style "
  "cross-device “pick and drop” image-transfer demonstration is built on "
  "top of the recognizer. All code, models and the collection tools for an "
  "open landmark-sequence dataset are released.")
p("Keywords: hand gesture recognition, early-exit networks, anytime "
  "inference, energy-efficient deep learning, proximity sensing, MediaPipe, "
  "on-device machine learning.", italic=True)

# ======================================================================= TOC
h1("Table of Contents")
toc_field()

# =========================================================== 1 INTRODUCTION
h1("1  Introduction")
h2("1.1  Motivation")
p("Hand gestures are a natural interaction modality when touching a device "
  "is inconvenient or impossible — cooking, driving, sterile environments, "
  "or accessibility contexts. Modern smartphones already contain everything "
  "needed to recognize gestures: a front camera and enough compute to run "
  "small neural networks. What they do not have is the energy budget to "
  "leave that pipeline running. A camera streaming at 15–30 fps together "
  "with per-frame model inference is among the most power-hungry workloads "
  "a phone can sustain, and an assistant that visibly drains the battery "
  "will simply be turned off. The central engineering question is therefore "
  "not whether gestures can be recognized, but how little energy a "
  "recognition system can consume while remaining responsive.")
h2("1.2  Research question and approach")
p("This thesis asks: how early in a hand gesture can an on-device model "
  "commit to a prediction, and how much energy does early commitment save — "
  "under a policy that can adapt to the phone's remaining battery? Our "
  "answer is architectural. First, the camera should not run while nothing "
  "is happening: a hardware proximity/light wave trigger (Tier 1) gates the "
  "entire vision pipeline. Second, once the camera is on, it should stop as "
  "soon as possible: a temporal early-exit network (Tier 2) evaluates "
  "intermediate classification heads as frames stream in and terminates the "
  "capture the moment it is confident. Both mechanisms attack the same "
  "quantity — camera-on time — which dominates the energy cost of the "
  "pipeline.")
h2("1.3  Contributions")
p("1. A temporal early-exit gesture model: a compact GRU over hand-landmark "
  "features with anytime-labeled exit heads at 4/8/12/16/24 frames and "
  "per-head temperature calibration, enabling accuracy-versus-latency "
  "operating points selected by a single threshold τ (Chapter 4).")
p("2. A battery-aware system design in which τ is the control knob of a "
  "battery-conditioned exit policy (designed in Chapter 3; measurement "
  "protocol in Chapter 6).")
p("3. A measured two-tier system on commodity phones: a sensor-fused wave "
  "trigger gating a camera + landmark pipeline, implemented as an Android "
  "application with cloud-logged instrumentation, evaluated end-to-end, and "
  "demonstrated with a cross-device image-transfer application (Chapters 5 "
  "and 6).")
p("We explicitly do not claim novelty for early-exit networks, early action "
  "prediction, wave-to-wake triggers, or energy-aware inference in "
  "isolation; these exist in prior work (Chapter 2). The claim is their "
  "specific combination, on-device, with an honest measurement methodology.")
h2("1.4  Scope and honesty of metrics")
p("Sensing is active only while the screen is on and the application is "
  "foregrounded, matching platform restrictions for third-party apps. "
  "Because no external power instrument is used, this thesis never reports "
  "absolute power in watts. The primary metrics are frames-to-decision, "
  "camera-on time and duty cycle, on-device latency, and relative battery "
  "drain per session measured by the operating system's battery gauge, "
  "reported as mean ± sd over repeated controlled sessions.")

# ====================================================== 2 LITERATURE REVIEW
h1("2  Background and Related Work")
h2("2.1  Early-exit and anytime inference")
p("Early-exit networks attach auxiliary classifiers to intermediate layers "
  "so that easy inputs can leave the network early. BranchyNet [3] "
  "introduced joint training of branch classifiers; MSDNet [4] developed "
  "multi-scale dense architectures with anytime prediction as an explicit "
  "objective. Surveys of adaptive inference [6] catalogue exit policies, "
  "training losses and calibration issues. In the video domain, FrameExit "
  "[5] performs per-frame early exiting for efficient video recognition — "
  "the nearest architectural relative of our Tier 2 — but targets server-"
  "side video understanding rather than on-device interaction, and has no "
  "notion of a battery-conditioned policy. Applications of early exits to "
  "wearable human-activity recognition exist [13], and a dynamic-inference "
  "surface-EMG gesture system on a microcontroller [2] is the closest prior "
  "in gesture specifically; we found no prior work applying temporal early "
  "exits to camera/landmark-based hand gesture recognition on phones.")
h2("2.2  Hand gesture recognition on mobile devices")
p("MediaPipe Hands [7] provides real-time 21-point hand landmark estimation "
  "on mobile hardware and has become the de facto front end for lightweight "
  "gesture systems; training compact temporal models on landmark sequences "
  "rather than pixels drastically reduces model size and domain gap. "
  "Large-scale gesture video datasets include Jester [9] and the phone-"
  "oriented IPN Hand [8]. Ultrasonic sensing offers a camera-free "
  "alternative: Fertl et al. [1] demonstrate end-to-end ultrasonic gesture "
  "recognition and explicitly list per-pulse early prediction as open "
  "future work, which supports the relevance of the early-commit question. "
  "SoundWave [10] showed Doppler-based gesture sensing using unmodified "
  "speaker and microphone, a candidate future Tier-1 alternative.")
h2("2.3  Energy-aware on-device inference")
p("Energy-aware scheduling of DNN inference under battery constraints has "
  "been studied for intermittent and edge systems (e.g., Zygarde [11]) and "
  "through self-adaptive approximate computing [12]. These works adapt "
  "which model or how much computation runs; our battery policy instead "
  "adapts within a single model, by scaling the exit threshold τ with the "
  "state of charge — full battery buys deliberation, low battery buys "
  "frugality. To our knowledge battery-conditioned exit thresholds for "
  "gesture recognition are unexplored.")
h2("2.4  Positioning")
p("Dusk Protocol combines (i) a zero-ML sensor wake trigger, (ii) landmark-"
  "based tiny temporal models, and (iii) anytime early exits with a "
  "battery-conditioned threshold, evaluated with honest energy proxies on "
  "commodity hardware. Each ingredient is established; the integration and "
  "its end-to-end measurement are the contribution.")

# ========================================================== 3 SYSTEM DESIGN
h1("3  System Design")
figure("fig_architecture.png",
       "Figure 3.1 — Dusk Protocol two-tier architecture. Tier 1 keeps the "
       "camera off until a deliberate wave is sensed; Tier 2 stops the "
       "camera at the first confident exit head.")
h2("3.1  Tier 1 — sensor-fused wave trigger")
p("Tier 1 must run continuously, so it must be effectively free. It is a "
  "rule-based state machine (~60 lines, no ML) over two event-driven "
  "hardware sensors. The infrared proximity sensor reports near/far "
  "transitions; three near-onsets within 2 s constitute a wave. Because "
  "many phones expose proximity only as a 1-bit signal with a range of a "
  "few centimetres, we fuse a second signal: the ambient light sensor. A "
  "hand waved 7–15 cm above the phone casts a shadow, producing "
  "characteristic dips in the lux reading; two dips of ≥22% below an "
  "exponentially tracked baseline within 3 s also fire the trigger. The "
  "light path requires ambient illumination (≥10 lux) and the proximity "
  "path covers darkness, so the fusion degrades gracefully. Sensor "
  "characteristics vary substantially across phone models; we report this "
  "variability as a finding rather than hiding it.")
h2("3.2  Tier 2 — camera, landmarks, and early-exit inference")
p("On waking, the front camera streams at the device rate (15–30 fps). "
  "Each frame passes through MediaPipe Hands [7] to produce 21 landmark "
  "coordinates. Frames are fed to the recognizer at ~15 fps to match the "
  "training frame rate. From each frame we compute 14 scale-normalized "
  "pairwise distances (fingertip–wrist openness, fingertip–knuckle curl, "
  "adjacent-fingertip spread), plus their first differences — 28 features "
  "that are invariant to translation, rotation, mirroring and camera "
  "distance (Section 4.3). The early-exit GRU consumes this stream; at "
  "frames 4, 8, 12, 16 and 24 an exit head produces a calibrated class "
  "distribution, and the first head whose confidence exceeds τ commits. On "
  "commit — or on a 10 s timeout — the camera is switched off (the media "
  "stream is torn down, not merely hidden), returning the system to Tier 1.")
h2("3.3  Battery-aware exit policy")
p("The exit threshold τ trades accuracy for frames: higher τ waits longer "
  "but errs less. Since the cost of waiting is camera-on time, τ is the "
  "natural control point for battery adaptation. The designed policy is a "
  "step rule — τ = 0.95 above 60% charge, 0.90 between 30–60%, and 0.80 "
  "below 30% — to be refined into an optimized mapping on a drain model "
  "calibrated with on-phone measurements. The application already exposes "
  "τ as a runtime parameter and logs battery state with every event, so "
  "the policy chapter reduces to an evaluation exercise on the operating "
  "curve of Figure 4.4.")
h2("3.4  Gesture vocabulary")
p("The prototype vocabulary is the transfer pair used by the demonstration "
  "application: grab (an open hand closing into a fist) and drop (a fist "
  "opening into a spread hand), with a third none class covering everything "
  "else — including deliberately adversarial near-misses such as half-open "
  "hands, counting poses, and thumbs-up. Restricting to two target "
  "gestures keeps the demonstrator crisp while exercising every part of "
  "the pipeline; the architecture is class-count agnostic.")

# ========================================================== 4 GESTURE MODELS
h1("4  Gesture Models")
h2("4.1  Static baseline: silhouette CNN")
p("Development began with a static hand-pose classifier to de-risk the "
  "pipeline. The training corpus is a 24,000-image silhouette dataset "
  "(50×50 binary masks, 20 pose classes, 900 train / 300 test per class), "
  "remapped to the 3-class vocabulary: the three fist classes form grab, "
  "the two spread-hand classes form drop, and the remaining fifteen poses "
  "— including visually adjacent ones such as four-finger spreads and "
  "thumbs-up — form none. Keeping near-misses inside none forces a tight "
  "decision boundary around the target gestures and is the main source of "
  "the model's precision.")
table(["3-class label", "Source silhouette classes", "Content"],
      [["grab", "11, 14, 16", "solid fists from three viewpoints"],
       ["drop", "4, 5", "fully spread open hands"],
       ["none", "remaining 15", "counting poses, OK, thumbs-up, bent hands…"]],
      "Table 4.1 — Dataset remapping from 20 pose classes to the working vocabulary.")
p("The classifier is a 287k-parameter CNN (batch-normalized convolutions "
  "with global average pooling) over 64×64 binarized inputs, trained with "
  "balanced sampling, label smoothing, and augmentation including random "
  "erosion/dilation of the silhouettes. It reaches 100.0% on the 6,000-"
  "image held-out test set.")
h2("4.2  Crossing the domain gap: rendering and synthetic hands")
p("A live camera provides RGB pixels, not segmentation masks. Rather than "
  "attempt fragile skin segmentation, we classify a silhouette rendered "
  "from the MediaPipe landmarks themselves: a filled palm polygon, thick "
  "finger capsules, and a forearm stub, followed by a morphological "
  "closing. The rendering is lighting- and skin-tone-independent by "
  "construction. To make the model robust in this rendered domain we built "
  "a procedural hand generator: a parametric skeleton (per-finger curl, "
  "spread, thumb pose) sampled into labeled poses by construction "
  "(Figure 4.1). Training jointly on both domains — selecting checkpoints "
  "by the worse of the two validation scores — yields 100.0% on the "
  "dataset test and 99.0% on 1,500 synthetic landmark renders "
  "(Figure 4.2).")
figure("fig_synth_poses.png",
       "Figure 4.1 — Procedurally generated hand poses rendered through the "
       "landmark-silhouette pipeline (rows: grab / drop / none).")
figure("fig_confusion_static.png",
       "Figure 4.2 — Static model confusion matrix on the landmark-render "
       "test domain (the domain live inference operates in).", 9.5)
h2("4.3  Temporal early-exit model")
h3("4.3.1  Features")
p("A grab is a motion, not a pose. The temporal model therefore consumes "
  "sequences of per-frame feature vectors: five fingertip–wrist distances "
  "(openness), five fingertip–MCP distances (curl), and four adjacent-"
  "fingertip gaps (spread), all normalized by the wrist–middle-MCP "
  "distance, plus first differences — 28 dimensions per frame, invariant "
  "to position, rotation, scale and handedness. This invariance is what "
  "lets synthetic training sequences and live MediaPipe streams share one "
  "input domain.")
h3("4.3.2  Synthetic sequences with anytime labels")
p("Training sequences are generated by animating the parametric skeleton: "
  "grab interpolates open→fist and drop fist→open with randomized onset "
  "(0–10 frames), duration (4–14 frames), per-finger lag and landmark "
  "jitter; none comprises hard negatives — held poses (including held "
  "fists and held open hands), half-grasps that re-open, finger-count "
  "changes, and thumb-only motion (Figure 4.3). Because the generator "
  "knows the gesture progress p(t), each exit head k receives an honest "
  "anytime label: none while p(t_k) ≤ 0.1, the gesture class once "
  "p(t_k) ≥ 0.7, and a masked (ignored) label in between. This prevents "
  "the pathological supervision that arises when early heads are forced to "
  "predict a gesture that has not yet begun.")
figure("fig_synth_seq.png",
       "Figure 4.3 — Synthetic sequences (every 2nd frame): grab, drop, and "
       "two none feints, with per-head anytime labels derived from "
       "generation-time progress.")
h3("4.3.3  Architecture, training and calibration")
p("The model is a single-layer GRU (hidden size 128; 62,607 parameters — "
  "well inside the 50–200k budget for real-time on-device inference) with "
  "a linear classification head at frames 4, 8, 12, 16 and 24. The anytime "
  "loss is a weighted sum of per-head cross-entropies (weights 0.5–1.0, "
  "masked labels excluded). After training, one softmax temperature per "
  "head is fitted on validation data so that a single τ is meaningful "
  "across heads. Training uses 12,000 fresh sequences per epoch for 20 "
  "epochs (≈7 minutes on CPU), and streaming inference is exactly "
  "equivalent to batch inference, verified to 1e-5.")
h3("4.3.4  Results")
figure("fig_head_acc.png",
       "Figure 4.4a — Accuracy at each exit head against anytime labels "
       "(synthetic test, n = 3,000).", 11)
tau_rows = [[r["tau"], r["gesture_acc"], r["mean_frames"], r["median_frames"],
             r["grab_frames"], r["drop_frames"], r["false_fire"]]
            for r in D["tau_rows"]]
table(["τ", "gesture acc.", "mean frames", "median", "grab frames",
       "drop frames", "false-fire"],
      tau_rows,
      "Table 4.2 — Exit-policy sweep on the synthetic test set "
      "(2,000 gesture / 1,000 none sequences).")
figure("fig_tau_sweep.png",
       "Figure 4.4 — Early-exit operating curve: mean frames-to-decision "
       "and false-fire rate versus τ. Gesture accuracy is 100% at every "
       "operating point; even at τ = 0.99 the system decides in a mean of "
       "13.6 of 24 frames — 43% below the no-early-exit baseline — with "
       "zero false fires.")
p("Two caveats frame these numbers honestly. First, they are measured on "
  "synthetic sequences; they validate the architecture, the anytime-label "
  "training scheme and the measurement harness rather than constituting "
  "field accuracy. Live operation on real hands (Chapter 6) and a recorded "
  "real-sequence corpus (the collection tool ships with the system) are "
  "the corresponding field evidence, with a fine-tuning path already "
  "implemented. Second, frames-to-decision includes the randomized gesture "
  "onset, so its floor is set by gesture physics — the model commits at "
  "the first head after the motion completes, which is precisely the "
  "designed anytime behaviour.")

# ========================================================= 5 IMPLEMENTATION
h1("5  Implementation")
h2("5.1  Web prototype")
p("A React web prototype established the demonstration scenario — the "
  "cross-device “pick and drop”: grabbing over phone A uploads a selected "
  "image; dropping over phone B retrieves it. An Express backend "
  "(deployed behind an nginx reverse proxy) exposes upload and drop "
  "endpoints, with transfer state and all measurement logs persisted in a "
  "cloud PostgreSQL database so that server restarts cannot lose state. "
  "The web prototype is retained as a data-collection and recruitment "
  "tool; it is not an evaluation platform.")
h2("5.2  Android application")
p("The Android application wraps the React client in a Capacitor WebView "
  "with two pieces of platform-specific engineering. First, a native "
  "sensor plugin (~70 lines of Java) streams proximity and ambient-light "
  "events to the JavaScript wave detectors. Second, the temporal GRU runs "
  "in ~60 lines of dependency-free JavaScript directly on exported "
  "weights: at 62k parameters an inference step costs ~60k multiply-"
  "accumulates, microseconds on any phone, making an ML runtime "
  "unnecessary. MediaPipe hand landmarks are computed by the official "
  "on-device WASM/GPU pipeline. Camera off is physical: the media stream "
  "is stopped, not hidden. The app ships two experiment arms in one "
  "binary — always-on (camera streams continuously) and Dusk (proximity-"
  "gated) — selected at launch, so A/B comparisons run identical code.")
h2("5.3  Instrumentation and data accumulation")
p("Every state transition — wake (with triggering sensor), camera on/off "
  "(with interval length), gesture commit (with class, confidence, "
  "frames-to-decision and wake-to-commit latency), ignored commit, "
  "timeout — plus a battery sample every 60 s, is queued locally and "
  "flushed in the background to the server, which stores it in the "
  "database. An analysis script reproduces every evaluation table and "
  "figure in this thesis directly from that database, making the "
  "measurement chain auditable end-to-end.")

# ============================================================= 6 EVALUATION
h1("6  Evaluation")
h2("6.1  Metrics")
table(["Metric", "Definition", "Source"],
      [["Camera duty cycle", "camera-on time / session time", "camera_on/off events"],
       ["Frames-to-decision", "frames consumed before commit", "commit events"],
       ["Wake→commit latency", "wave trigger to gesture commit", "commit events"],
       ["Battery drain", "battery %/30 min per mode (mean ± sd)", "60 s battery samples"],
       ["Tier-1 wakes / timeouts", "trigger counts; timeout = wake without gesture", "wake/timeout events"]],
      "Table 6.1 — Evaluation metrics. No absolute power (watts) is claimed anywhere in this thesis.")
h2("6.2  Measured camera duty cycle")
p(f"Across {D['n_always']} always-on and {D['n_dusk']} Dusk live sessions "
  "logged to date, the always-on arm kept the camera streaming for "
  f"{D['duty_always']:.1f}% of session time, while the proximity-gated arm "
  f"required only {D['duty_dusk']:.1f}% — a "
  f"{D['duty_always']/max(D['duty_dusk'],1e-9):.1f}× reduction in the "
  "quantity that dominates the pipeline's energy cost (Figure 6.1). "
  "Measured wake-to-commit latency in the Dusk arm ranged 3.2–7.3 s, "
  "comprising the wave, camera start-up, hand acquisition and the "
  "gesture itself.")
figure("fig_camera_duty.png",
       "Figure 6.1 — Measured camera duty cycle per experiment arm "
       "(live sessions, identical binary).", 11)
h2("6.3  Battery: simulated projection and measurement protocol")
p("A controlled battery A/B (30-minute scripted sessions, fixed brightness, "
  "airplane mode with Wi-Fi, one interaction every 2 minutes, ≥3 repeats "
  "per arm in ABBA order on the same handset) is specified in the project's "
  "measurement protocol and is being executed at the time of this draft. "
  "To bound the expected effect, Figure 6.2 presents an explicitly "
  "SIMULATED projection: battery trajectories computed from the measured "
  "duty cycles of Section 6.2 under an assumed screen-on baseline of "
  "4%/30 min and a camera-pipeline cost κ swept over 4–8%/30 min at 100% "
  "duty. Under the mid assumption (κ = 6), the always-on arm drains "
  f"{D['proj_always']:.1f}%/30 min against {D['proj_dusk']:.1f}%/30 min "
  f"for Dusk — a {D['proj_saving']:.0f}% reduction. These numbers are a "
  "model, not a measurement; they will be replaced by the measured A/B "
  "in the final thesis.")
figure("fig_battery_projection.png",
       "Figure 6.2 — SIMULATED battery projection from measured duty "
       "cycles (shaded: κ sensitivity band). To be replaced by the "
       "measured on-phone A/B.")
h2("6.4  Tier-1 behaviour")
p("On the primary test handset (Infinix X6812), the IR proximity sensor "
  "behaves as a binary near/far switch with an effective range of at most "
  "a few centimetres, motivating the light-sensor fusion of Section 3.1; "
  "wake events from both sources are recorded in the log with their "
  "triggering sensor. Timeout events (wake without a subsequent gesture) "
  "serve as the false-wake proxy and are part of the standing protocol, "
  "including a no-intended-use session to measure spurious triggers "
  "during ordinary desk handling.")
h2("6.5  Threats to validity")
p("(i) The temporal model's headline accuracies are on synthetic sequences; "
  "real-sequence fine-tuning and evaluation are implemented but the "
  "recorded corpus is still small. (ii) Battery results in this draft are "
  "simulation-labeled projections pending the controlled A/B. (iii) Duty-"
  "cycle sessions to date are exploratory rather than scripted, and some "
  "were collected on different devices; the protocol mandates same-device "
  "comparisons. (iv) Sensor characteristics (proximity range, light-sensor "
  "update rate) vary across handsets; results are reported per device.")

# ================================================== 7 LIMITATIONS & FUTURE
h1("7  Limitations and Future Work")
p("Beyond the threats above: the system senses only while the app is "
  "foregrounded with the screen on, by platform policy; the wave trigger's "
  "light path needs ambient illumination; and the two-gesture vocabulary, "
  "while sufficient for the transfer demonstration, does not exercise "
  "per-class exit-behaviour analysis at scale. Future work includes: the "
  "measured battery A/B and the calibrated battery-conditioned τ policy; "
  "fine-tuning and evaluation on the recorded real-sequence corpus and on "
  "IPN Hand [8]; per-class frames-to-decision analysis (“which gestures "
  "are easy?”); an ultrasonic Doppler Tier-1 variant in the spirit of "
  "SoundWave [10] for longer-range waking; a native TFLite deployment for "
  "comparison with the in-WebView implementation; and the public release "
  "of the landmark-sequence dataset collected with the shipped tools.")

# ============================================================= 8 CONCLUSION
h1("8  Conclusion")
p("Dusk Protocol demonstrates that energy-frugal, always-available gesture "
  "interaction is achievable on unmodified smartphones by attacking "
  "camera-on time from two directions: do not start the camera until a "
  "free sensor says a hand is present, and stop it the moment a calibrated "
  "early-exit head is confident. The temporal early-exit recognizer "
  "reaches ceiling accuracy on its test distribution while cutting frames-"
  "to-decision by 43%, and the deployed two-tier system cuts measured "
  "camera duty cycle by a factor of five in live use. The full stack — "
  "models, training and synthesis code, the Android application, the "
  "logging backend and the analysis pipeline — is reproducible from a "
  "single repository, and every number in the final thesis will be "
  "regenerable from the project database.")

# =============================================================== REFERENCES
h1("References")
refs = [
 "M. Fertl, D. Palmieri, S. Hussain, et al., “End-to-End Ultrasonic Hand "
 "Gesture Recognition,” Sensors, vol. 24, no. 9, p. 2740, 2024.",
 "X. Xie et al., “Dynamic-inference surface-EMG gesture recognition on "
 "ultra-low-power edge processors,” Sensors, 2023.",
 "S. Teerapittayanon, B. McDanel, and H. T. Kung, “BranchyNet: Fast "
 "inference via early exiting from deep neural networks,” in Proc. ICPR, "
 "2016.",
 "G. Huang, D. Chen, T. Li, F. Wu, L. van der Maaten, and K. Weinberger, "
 "“Multi-scale dense networks for resource efficient image "
 "classification,” in Proc. ICLR, 2018.",
 "A. Ghodrati, B. E. Bejnordi, and A. Habibian, “FrameExit: Conditional "
 "early exiting for efficient video recognition,” in Proc. CVPR, 2021.",
 "S. Laskaridis, A. Kouris, and N. D. Lane, “Adaptive inference through "
 "early-exit networks: Design, challenges and directions,” in Proc. "
 "EMDL Workshop (MobiSys), 2021.",
 "F. Zhang, V. Bazarevsky, A. Vakunov, et al., “MediaPipe Hands: "
 "On-device real-time hand tracking,” arXiv:2006.10214, 2020.",
 "G. Benitez-Garcia, J. Olivares-Mercado, G. Sanchez-Perez, and K. Yanai, "
 "“IPN Hand: A video dataset and benchmark for real-time continuous hand "
 "gesture recognition,” in Proc. ICPR, 2020.",
 "J. Materzynska, G. Berger, I. Bax, and R. Memisevic, “The Jester "
 "dataset: A large-scale video dataset of human gestures,” in Proc. ICCV "
 "Workshops, 2019.",
 "S. Gupta, D. Morris, S. Patel, and D. Tan, “SoundWave: Using the "
 "Doppler effect to sense gestures,” in Proc. CHI, 2012.",
 "B. Islam and S. Nirjon, “Zygarde: Time-sensitive on-device deep "
 "inference and adaptation on intermittently-powered systems,” Proc. ACM "
 "IMWUT, vol. 4, no. 3, 2020.",
 "M. Tundo et al., “An energy-aware approach to design self-adaptive "
 "AI-based applications on the edge,” in Proc. ASE, 2023.",
 "E. Lattanzi et al., “Exploiting concurrency and early exits for "
 "energy-efficient human activity recognition,” Engineering Applications "
 "of Artificial Intelligence, 2023.",
]
for i, r in enumerate(refs, 1):
    para = p(f"[{i}] {r}", size=11)
    para.paragraph_format.space_after = Pt(4)
p("Note: per the project's citation policy, every reference above is to be "
  "re-verified against its DOI/arXiv page before final submission.",
  italic=True, size=10)

# ================================================================= APPENDIX
h1("Appendix A — Repository and Reproducibility")
p("All code is in a single repository: model training and synthesis "
  "(handgesture/), the Android application (android-prototype/), the web "
  "prototype and backend (web-prototype/), and the measurement protocol "
  "and analysis (measurements/). Key entry points:")
table(["Artifact", "Command / location"],
      [["Static model training", "handgesture/train3.py"],
       ["Temporal early-exit training + τ sweep", "handgesture/train_temporal.py"],
       ["Real-sequence recorder (dataset tool)", "handgesture/record_sequences.py"],
       ["Model export for the app", "handgesture/export_gru_json.py"],
       ["Android APK build", "android-prototype: npm run apk"],
       ["Battery A/B protocol", "measurements/PROTOCOL.md"],
       ["Regenerate all evaluation figures", "measurements/analyze_logs.py"],
       ["Regenerate this document", "docs/build_thesis.py"]],
      "Table A.1 — Reproducibility map.")

doc.save(OUT)
print("saved", OUT)
